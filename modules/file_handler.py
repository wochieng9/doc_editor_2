
"""
File Handling Module
------------------
Manages local file operations for the Streamlit Document Editor,
including file reading, writing, and format conversions.

Author: Walter Ochieng
Email: ocu9@cdc.gov
Version: 1.0.0
Date: 2025-01-21
License: MIT

Supported File Formats:
    - Plain Text (.txt)
    - Markdown (.md)
    - Rich Text (.rtf)
    - Word Documents (.docx)

Dependencies:
    - python-docx>=0.8.11

Usage:
    from modules.file_handling import FileHandler
    
    handler = FileHandler()
    content = handler.read_file('document.docx')
    handler.save_file('output.txt', content)
"""

from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import PyPDF2
import docx
import io
from typing import List, Dict, Union
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
class FileHandler:

    def read_file(self, file):
        """Read text from uploaded file"""
        try:
            file_type = file.name.split('.')[-1].lower()
            
            if file_type == 'txt':
                return file.getvalue().decode('utf-8')
            elif file_type == 'docx':
                doc = Document(file)
                return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            elif file_type == 'pdf':
                pdf_reader = PyPDF2.PdfReader(file)
                return '\n'.join(page.extract_text() for page in pdf_reader.pages)
            return None
        except Exception as e:
            logger.error(f"Error in reading file: {str(e)}")
            raise

    
    def validate_changes(self, changes: List[Dict[str, Union[int, str]]]) -> bool:
        """
        Validate the changes list structure
        Returns True if valid, raises ValueError if invalid
        """
        if not isinstance(changes, list):
            raise ValueError(f"Changes must be a list, got {type(changes)}")
        for i, change in enumerate(changes):
            if not isinstance(change, dict):
                raise ValueError(f"Change at index {i} must be a dictionary, got {type(change)}")
            if "position" not in change or "text" not in change:
                raise ValueError(f"Change at index{i} missing required keys 'position' and 'text'")
            if not isinstance(change["position"], int):
                raise ValueError(f"Position at index {i} must be an integer, got {type(change['position'])}")
            if not isinstance(change["text"], str):
                raise ValueError(f"Text at index {i} must be a string, got {type(change['text'])}")

        return True

    def convert_changes_format(self, changes: List) -> List[Dict[str, Union[int, str]]]:
        if not changes:
            return []
        if isinstance(changes[0], dict) and all(key in changes[0] for key in ["position", "text"]):
            return changes
        converted_changes = []
        for change in changes:
            if isinstance(change, List) and len(change) >= 3:
                _, position, text = change
                converted_changes.append({
                    "position": position,
                    "text": text,
                })
            else:
                logger.warning(f"Skipping invalid change format: {change}")
        return converted_changes

                        

    def save_edited_document(self, text: str, format_type: str, changes: List = None) -> io.BytesIO:
        """
        Save document in a specified format with highlighted changes
        Args:
            text: Content to save
            format_type: Output format ('docx' or 'pdf')
            changes: List of changes (either from highlight_differences or in dictionary format)

        Returns:
            io.BytesIO: Document in specified format
        """
        try:
            logger.info(f"Received changes: {changes}")

            if changes:
                formatted_changes = self.convert_changes_format(changes)
                try:
                    self.validate_changes(formatted_changes)
                except ValueError as e:
                    logger.error(f"Invalid changes format: {str(e)}")
                    raise
                changes = formatted_changes

            if format_type == "docx":
                doc = Document()
                current_paragraph = None
                current_pos = 0

                if changes:
                    sorted_changes = sorted(changes, key=lambda x: x['position'])
                    
                    for change in sorted_changes:
                        # Add unchanged text before the change
                        if current_pos < change['position']:
                            text_segment = text[current_pos:change['position']]
                            # Only split if there's an actual newline in the text
                            segments = text_segment.split('\n')
                            
                            for i, segment in enumerate(segments):
                                if segment or i < len(segments) - 1:  # Add empty paragraphs only if they're explicit newlines
                                    if current_paragraph is None or i > 0:
                                        current_paragraph = doc.add_paragraph()
                                    if segment:
                                        current_paragraph.add_run(segment)
                        
                        # Add the changed text
                        changed_segments = change['text'].split('\n')
                        for i, segment in enumerate(changed_segments):
                            if segment or i < len(changed_segments) - 1:  # Add empty paragraphs only if they're explicit newlines
                                if current_paragraph is None or i > 0:
                                    current_paragraph = doc.add_paragraph()
                                if segment:
                                    run = current_paragraph.add_run(segment)
                                    run.font.color.rgb = docx.shared.RGBColor(255, 0, 0)
                        
                        current_pos = change['position'] + len(change['text'])
                    
                    # Add remaining text
                    if current_pos < len(text):
                        remaining_text = text[current_pos:]
                        segments = remaining_text.split('\n')
                        for i, segment in enumerate(segments):
                            if segment or i < len(segments) - 1:  # Add empty paragraphs only if they're explicit newlines
                                if current_paragraph is None or i > 0:
                                    current_paragraph = doc.add_paragraph()
                                if segment:
                                    current_paragraph.add_run(segment)
                else:
                    # No changes - just add the text normally
                    segments = text.split('\n')
                    for i, segment in enumerate(segments):
                        if segment or i < len(segments) - 1:  # Add empty paragraphs only if they're explicit newlines
                            current_paragraph = doc.add_paragraph()
                            if segment:
                                current_paragraph.add_run(segment)

                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                return doc_io

            elif format_type == 'pdf':
                pdf_io = io.BytesIO()
                c = canvas.Canvas(pdf_io, pagesize=letter)
                text_object = c.beginText(40, 750)
                text_object.setFont("Helvetica", 12)
                
                if changes:
                    current_pos = 0
                    sorted_changes = sorted(changes, key=lambda x: x['position'])
                    
                    for change in sorted_changes:
                        # Add unchanged text
                        if current_pos < change['position']:
                            text_segment = text[current_pos:change['position']]
                            text_object.setFillColor('black')
                            lines = text_segment.rstrip('\n').split('\n')
                            for line in lines:
                                text_object.textLine(line)
                        
                        # Add changed text
                        text_object.setFillColor('red')
                        changed_lines = change['text'].rstrip('\n').split('\n')
                        for line in changed_lines:
                            text_object.textLine(line)
                                
                        current_pos = change['position'] + len(change['text'])
                    
                    # Add remaining text
                    if current_pos < len(text):
                        text_object.setFillColor('black')
                        remaining_lines = text[current_pos:].rstrip('\n').split('\n')
                        for line in remaining_lines:
                            text_object.textLine(line)
                else:
                    # No changes - just add the text normally
                    lines = text.rstrip('\n').split('\n')
                    for line in lines:
                        text_object.textLine(line)

                c.drawText(text_object)
                c.save()
                pdf_io.seek(0)
                return pdf_io
            else:
                raise ValueError(f"Unsupported format: {format_type}")
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            raise



















                        
        
                                     

